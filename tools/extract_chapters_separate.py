from docx import Document
import os

chapters = [
    ('Chapter_1_Introduction_EXPANDED_FULL.docx', 'chapter_1.txt'),
    ('Chapter2_Complete_Literature_Review.docx', 'chapter_2.txt'),
    ('Chapter3_Methodology.docx', 'chapter_3.txt')
]

base_path = r'F:\My Drive\Multimodal_Cancer_Detection\thesis\Dissertation'
output_path = r'F:\My Drive\Multimodal_Cancer_Detection'

for docx_file, txt_file in chapters:
    input_path = os.path.join(base_path, docx_file)
    output_file = os.path.join(output_path, txt_file)
    
    try:
        doc = Document(input_path)
        
        # Extract all text from paragraphs and tables
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        # Write to file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"{'='*80}\n")
            f.write(f"{docx_file}\n")
            f.write(f"{'='*80}\n\n")
            f.write('\n'.join(full_text))
        
        print(f"✓ Extracted {docx_file} to {txt_file} ({len(full_text)} paragraphs)")
        
    except Exception as e:
        print(f"✗ Error processing {docx_file}: {e}")

print("\nExtraction complete!")
