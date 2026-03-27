from __future__ import annotations

import os
from pathlib import Path

from docx import Document


def extract_docx_to_text(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if text:
                    parts.append(text)

    return parts


def main() -> None:
    base_dir = Path(r"F:\My Drive\Multimodal_Cancer_Detection\thesis\Dissertation_writing_instruction")
    out_file = Path(r"F:\My Drive\Multimodal_Cancer_Detection\dissertation_writing_instruction.txt")

    docx_files = sorted(base_dir.glob("*.docx"), key=lambda p: p.name.lower())
    if not docx_files:
        raise SystemExit(f"No .docx files found in {base_dir}")

    all_lines: list[str] = []
    for docx_path in docx_files:
        try:
            lines = extract_docx_to_text(docx_path)
        except Exception as exc:  # noqa: BLE001
            all_lines.append("=" * 80)
            all_lines.append(docx_path.name)
            all_lines.append("=" * 80)
            all_lines.append(f"[ERROR extracting this document: {exc}]")
            all_lines.append("")
            continue

        all_lines.append("=" * 80)
        all_lines.append(docx_path.name)
        all_lines.append("=" * 80)
        all_lines.extend(lines)
        all_lines.append("")

    out_file.write_text("\n".join(all_lines), encoding="utf-8")
    print(f"Wrote: {out_file} (from {len(docx_files)} .docx files)")


if __name__ == "__main__":
    main()
